from datetime import datetime


def generate_audit_docx(report, output_path: str):
    """
    Build a formatted Word document for the approved audit report.

    Args:
        report: AuditReport model instance
        output_path: absolute path to write the .docx file
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    GREEN  = RGBColor(0x06, 0x5F, 0x33)   # --g700
    DKGREY = RGBColor(0x1E, 0x29, 0x3B)   # slate-800
    LTGREY = RGBColor(0x94, 0xA3, 0xB8)   # slate-400
    WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

    doc = Document()

    # ── Page margins ──────────────────────────────────────
    section = doc.sections[0]
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)

    # ── Helper: set cell background ───────────────────────
    def shade_cell(cell, hex_color: str):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def set_cell_border(cell, **kwargs):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for edge in ("top", "bottom", "left", "right"):
            if edge in kwargs:
                tag = OxmlElement(f"w:{edge}")
                tag.set(qn("w:val"),   kwargs[edge].get("val",   "single"))
                tag.set(qn("w:sz"),    kwargs[edge].get("sz",    "4"))
                tag.set(qn("w:space"), "0")
                tag.set(qn("w:color"), kwargs[edge].get("color", "auto"))
                tcBorders.append(tag)
        tcPr.append(tcBorders)

    # ══════════════════════════════════════════════════════
    # COVER HEADER
    # ══════════════════════════════════════════════════════
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("CLUBS & ASSOCIATIONS MANAGEMENT SYSTEM")
    run.bold      = True
    run.font.size = Pt(16)
    run.font.color.rgb = GREEN

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Official Audit Report")
    sr.font.size  = Pt(12)
    sr.font.color.rgb = DKGREY

    doc.add_paragraph()  # spacer

    # ── Divider line ──────────────────────────────────────
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "16A34A")
    pBdr.append(bottom)
    pPr.append(pBdr)

    doc.add_paragraph()

    # ── Report identity table ─────────────────────────────
    id_table = doc.add_table(rows=3, cols=2)
    id_table.style = "Table Grid"
    id_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    def id_row(row_idx, label, value):
        row = id_table.rows[row_idx]
        shade_cell(row.cells[0], "F0FDF4")
        lc = row.cells[0].paragraphs[0]
        lc.add_run(label).bold = True
        vc = row.cells[1].paragraphs[0]
        vc.add_run(str(value))

    id_row(0, "Club Name",      report.club.name)
    id_row(1, "Audit Period",   f"{report.period.value} {report.year}")
    id_row(2, "Report Status",  "APPROVED ✓")

    doc.add_paragraph()

    meta_table = doc.add_table(rows=3, cols=2)
    meta_table.style = "Table Grid"
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    def meta_row(row_idx, label, value):
        row = meta_table.rows[row_idx]
        shade_cell(row.cells[0], "F0FDF4")
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(str(value))

    meta_row(0, "Submitted By",    report.submitter.full_name)
    meta_row(1, "Reviewed By",     report.reviewer.full_name if report.reviewer else "—")
    meta_row(2, "Date Approved",
             report.reviewed_at.strftime("%d %B %Y, %H:%M UTC") if report.reviewed_at else "—")

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════
    # SECTION HELPER
    # ══════════════════════════════════════════════════════
    def add_section_heading(text: str):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = WHITE
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "065F33")
        pPr.append(shd)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.left_indent  = Pt(6)

    def add_two_col_table(data: list[tuple], header=False):
        """data = [(label, value), ...]"""
        t = doc.add_table(rows=len(data), cols=2)
        t.style = "Table Grid"
        for i, (label, value) in enumerate(data):
            row = t.rows[i]
            shade_cell(row.cells[0], "F0FDF4")
            row.cells[0].paragraphs[0].add_run(str(label)).bold = True
            row.cells[1].paragraphs[0].add_run(str(value))
        doc.add_paragraph()

    # ══════════════════════════════════════════════════════
    # 1. MEMBERSHIP SUMMARY
    # ══════════════════════════════════════════════════════
    add_section_heading("1. Membership Summary")
    add_two_col_table([
        ("Total Members",           report.total_members),
        ("Active Members",          report.active_members),
        ("New Members This Period",  report.new_members),
        ("Members Who Left",         report.members_left),
    ])

    # ══════════════════════════════════════════════════════
    # 2. EVENTS & ACTIVITIES
    # ══════════════════════════════════════════════════════
    add_section_heading("2. Events & Activities")
    add_two_col_table([
        ("Events Held",             report.events_held),
        ("Events Planned Next Period", report.events_planned),
        ("Average Attendance",      f"{report.average_attendance:.1f} members/event"),
    ])

    # ══════════════════════════════════════════════════════
    # 3. FINANCIAL SUMMARY
    # ══════════════════════════════════════════════════════
    add_section_heading("3. Financial Summary")
    add_two_col_table([
        ("Opening Balance",    f"KES {report.opening_balance:,.2f}"),
        ("Total Income",       f"KES {report.total_income:,.2f}"),
        ("Total Expenditure",  f"KES {report.total_expenditure:,.2f}"),
        ("Closing Balance",    f"KES {report.closing_balance:,.2f}"),
        ("Net Balance",        f"KES {report.net_balance:,.2f}"),
        ("Fees Collected",     f"KES {report.fees_collected:,.2f}"),
        ("Outstanding Fees",   f"KES {report.outstanding_fees:,.2f}"),
    ])

    # ══════════════════════════════════════════════════════
    # 4. COMPLIANCE CHECKLIST
    # ══════════════════════════════════════════════════════
    add_section_heading("4. Compliance Checklist")
    t = doc.add_table(rows=len(report.compliance_checklist) + 1, cols=3)
    t.style = "Table Grid"

    # Header row
    hdr = t.rows[0]
    shade_cell(hdr.cells[0], "065F33")
    shade_cell(hdr.cells[1], "065F33")
    shade_cell(hdr.cells[2], "065F33")
    for cell, text in zip(hdr.cells, ["#", "Compliance Item", "Status"]):
        r = cell.paragraphs[0].add_run(text)
        r.bold = True
        r.font.color.rgb = WHITE

    for i, (item, passed) in enumerate(report.compliance_checklist.items(), start=1):
        row = t.rows[i]
        shade_cell(row.cells[0], "F0FDF4")
        row.cells[0].paragraphs[0].add_run(str(i))
        row.cells[1].paragraphs[0].add_run(item)
        status_run = row.cells[2].paragraphs[0].add_run("✓ Yes" if passed else "✗ No")
        status_run.bold = True
        status_run.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A) if passed else RGBColor(0xEF, 0x44, 0x44)

    score_p = doc.add_paragraph()
    score_p.paragraph_format.space_before = Pt(4)
    sr = score_p.add_run(f"Compliance Score: {report.compliance_score}%")
    sr.bold = True
    sr.font.color.rgb = GREEN if report.compliance_score >= 80 else RGBColor(0xEF, 0x44, 0x44)
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════
    # 5. NARRATIVE SECTIONS
    # ══════════════════════════════════════════════════════
    for heading, content in [
        ("5. Achievements",     report.achievements),
        ("6. Challenges",       report.challenges),
        ("7. Recommendations",  report.recommendations),
        ("8. Additional Notes", report.additional_notes),
    ]:
        add_section_heading(heading)
        body_p = doc.add_paragraph(content or "None reported.")
        body_p.paragraph_format.left_indent = Pt(6)
        doc.add_paragraph()

    # ══════════════════════════════════════════════════════
    # REVIEW NOTE (from dean)
    # ══════════════════════════════════════════════════════
    if report.review_note:
        add_section_heading("Dean's Review Notes")
        rn = doc.add_paragraph(report.review_note)
        rn.paragraph_format.left_indent = Pt(6)
        doc.add_paragraph()

    # ══════════════════════════════════════════════════════
    # SIGNATURE BLOCK
    # ══════════════════════════════════════════════════════
    doc.add_paragraph()
    sig_table = doc.add_table(rows=4, cols=2)
    sig_table.style = "Table Grid"

    sig_rows = [
        ("Club Leader",  report.submitter.full_name),
        ("Dean",         report.reviewer.full_name if report.reviewer else "—"),
        ("Date",         report.reviewed_at.strftime("%d %B %Y") if report.reviewed_at else "—"),
        ("Report Ref",   f"CAMS-AUDIT-{report.id:05d}"),
    ]
    for i, (label, value) in enumerate(sig_rows):
        row = sig_table.rows[i]
        shade_cell(row.cells[0], "F0FDF4")
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(value)

    doc.add_paragraph()

    # ── Footer note ───────────────────────────────────────
    footer_p = doc.add_paragraph(
        f"Generated automatically by CAMS on "
        f"{datetime.now().strftime('%d %B %Y at %H:%M UTC')}. "
        f"This document is official and should be retained for institutional records."
    )
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = LTGREY

    doc.save(output_path)